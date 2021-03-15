#encoding=utf-8
'''
pip install python-docx
'''
from docx import Document
import os
import json

#from PIL import Image,ImageDraw
#from io import BytesIO
'420800196912232170'
user=str(input("输入身份证号：")).strip()
qtype=str(input("输入题型（例如A3）：")).strip()
if not os.path.exists(user):
    print("不存在该身份证的爬取记录")
    exit(0)
document = Document("shiti.docx") #新建文档
cnt = 1
first_p = True

def addQuestion(doc,dat):
    global first_p
    if '多选' in dat['type']:
        dat['type']='(多选题)'
    if first_p:
        document.paragraphs[0].add_run(dat['type'].replace("(",u"【").replace(")",u"】").replace(" ",''))
        first_p=False
    else:
        p = document.add_paragraph(dat['type'].replace("(",u"【").replace(")",u"】").replace(" ",''))  # 添加一个段落
    p = document.add_paragraph('难度：简单')
    p = document.add_paragraph("题干：" + '(%s)%s'%(qtype,dat['title']))  # 添加一个段落

    if dat['pics']:
        p = document.add_paragraph()  # 添加一个段落
        r = p.add_run()  # 添加一个游程
        r.add_picture(dat['pics'])  # 在当前游程中插入图片

    if u'判断' not in dat['type']:
        p = document.add_paragraph('选项：')      
        for op in sorted(dat['options']):
            if not dat['options'][op].startswith(op+"、"):
                dat['options'][op]=op+"、"+ dat['options'][op][1:]
            p = document.add_paragraph(dat['options'][op].replace("、","."))  # 添加一个段落

    answers =[]
    for x in dat['answer']:
        answers.append(x)
    if u'判断' in dat['type']:
        s =''
        answers = []
        for x in dat['answer']:
            if x == 'A':
                answers.append("对")
            if x == 'B':
                answers.append("错")
    p = document.add_paragraph('答案：%s' % ('，'.join(answers).replace("，",'')))
    p = document.add_paragraph('解析：')
    p = document.add_paragraph('【结束】')





files=[]
for filename in os.listdir(user):
    if filename.endswith(".json") and filename!= "info.json":
        files.append(int(filename[:-5]))

a={}
files.sort()
for filename in files:
       filename = "%d.json"%filename
       dat =json.load(open(user+'/'+filename))
       if dat['title'] not in a:
           a[dat['title']]=[]
       a[dat['title']].append(filename)

files.sort()
for filename in files:
       filename = "%d.json"%filename
       dat =json.load(open(user+'/'+filename))
       if len(a[dat['title']])>1:
           for i in range(1,len(a[dat['title']])+1):
               if a[dat['title']][i-1] == filename:
                   dat['title']=("%d-" % i) + dat['title']
                   print(dat)
                   break
       addQuestion(document,dat)


p = document.add_paragraph("")  # 添加最后的huanhang
document.add_paragraph("")  # 添加最后的huanhang
info = json.load(open("%s/info.json" % user))
document.save("%s_%s.docx"%(qtype,info['project']))